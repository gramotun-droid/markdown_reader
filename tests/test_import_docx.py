from __future__ import annotations

from pathlib import Path

import pytest

pytest.importorskip("docx")

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from app.import_docx import DocxImportError, docx_to_markdown  # noqa: E402


def _save(document, tmp_path: Path) -> Path:
    dest = tmp_path / "sample.docx"
    document.save(str(dest))
    return dest


def test_headings_and_paragraphs(tmp_path: Path):
    d = Document()
    d.add_heading("Title here", level=0)  # style "Title" -> level 1
    d.add_paragraph("A plain paragraph.")
    d.add_heading("Sub", level=2)
    md = docx_to_markdown(_save(d, tmp_path))
    assert "# Title here" in md
    assert "## Sub" in md
    assert "A plain paragraph." in md


def test_bold_italic_runs(tmp_path: Path):
    d = Document()
    p = d.add_paragraph()
    p.add_run("normal ")
    p.add_run("b").bold = True
    p.add_run(" ")
    p.add_run("i").italic = True
    both = p.add_run("x")
    both.bold = True
    both.italic = True
    md = docx_to_markdown(_save(d, tmp_path))
    assert "**b**" in md
    assert "*i*" in md
    assert "***x***" in md


def test_bullet_and_numbered_lists_stay_separate(tmp_path: Path):
    d = Document()
    d.add_paragraph("one", style="List Bullet")
    d.add_paragraph("two", style="List Bullet")
    d.add_paragraph("first", style="List Number")
    md = docx_to_markdown(_save(d, tmp_path))
    assert "- one" in md
    assert "- two" in md
    assert "1. first" in md
    # a blank line must separate the two lists so they render as distinct lists
    assert "- two\n\n1. first" in md


def test_quote_and_table(tmp_path: Path):
    d = Document()
    d.add_paragraph("quoted", style="Intense Quote")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "A"
    t.cell(0, 1).text = "B"
    t.cell(1, 0).text = "1"
    t.cell(1, 1).text = "2"
    md = docx_to_markdown(_save(d, tmp_path))
    assert "> quoted" in md
    assert "| A | B |" in md
    assert "| --- | --- |" in md
    assert "| 1 | 2 |" in md


def test_hyperlink(tmp_path: Path):
    d = Document()
    para = d.add_paragraph("see ")
    rid = para.part.relate_to(
        "https://example.com",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = para._p.makeelement(qn("w:hyperlink"), {qn("r:id"): rid})
    run = para._p.makeelement(qn("w:r"), {})
    text = run.makeelement(qn("w:t"), {})
    text.text = "site"
    run.append(text)
    hyperlink.append(run)
    para._p.append(hyperlink)
    md = docx_to_markdown(_save(d, tmp_path))
    assert "[site](https://example.com)" in md


def test_special_chars_are_escaped(tmp_path: Path):
    d = Document()
    d.add_paragraph("a_b*c[d]")
    md = docx_to_markdown(_save(d, tmp_path))
    assert r"a\_b\*c\[d\]" in md


def test_missing_file_raises(tmp_path: Path):
    with pytest.raises(DocxImportError):
        docx_to_markdown(tmp_path / "nope.docx")
