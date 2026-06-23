"""Export the open Markdown document to other formats.

PDF export is handled by the Qt WebEngine page (it prints the rendered page,
so the result matches what the user sees). DOCX export is built here by
walking the markdown-it token stream and emitting a Word document with
python-docx, covering the common Markdown constructs (headings, emphasis,
lists, quotes, code blocks, tables, rules and links).
"""

from __future__ import annotations

from pathlib import Path

from markdown_it import MarkdownIt
from mdit_py_plugins.tasklists import tasklists_plugin


class ExportError(RuntimeError):
    """Raised when a document cannot be exported."""


# Mapping of nesting depth -> python-docx built-in list styles.
_MAX_LIST_STYLE_LEVEL = 3


def _build_parser() -> MarkdownIt:
    return (
        MarkdownIt("commonmark", {"html": False, "linkify": True, "typographer": True})
        .enable("table")
        .enable("strikethrough")
        .use(tasklists_plugin, enabled=True)
    )


def markdown_to_docx(markdown_text: str, dest: Path, title: str = "") -> None:
    """Convert ``markdown_text`` to a .docx file at ``dest``."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - only when packaging is broken
        raise ExportError(
            "Для экспорта в DOCX требуется пакет python-docx, который не найден."
        ) from exc

    try:
        tokens = _build_parser().parse(markdown_text)
        document = Document()
        if title:
            document.core_properties.title = title
        _render_blocks(document, tokens)
        document.save(str(dest))
    except ExportError:
        raise
    except Exception as exc:  # noqa: BLE001 - surface any failure to the UI
        raise ExportError(f"Не удалось создать DOCX: {exc}") from exc


# --------------------------------------------------------------------- blocks


def _render_blocks(document, tokens) -> None:
    list_stack: list[str] = []  # "bullet" / "number" per open list level
    in_quote = 0
    index = 0
    count = len(tokens)

    while index < count:
        token = tokens[index]
        kind = token.type

        if kind == "heading_open":
            level = int(token.tag[1])
            paragraph = document.add_heading(level=min(level, 9))
            _render_inline(paragraph, tokens[index + 1])
            index += 3
            continue

        if kind == "paragraph_open":
            paragraph = _new_paragraph(document, list_stack, in_quote)
            _render_inline(paragraph, tokens[index + 1])
            index += 3
            continue

        if kind in ("bullet_list_open", "ordered_list_open"):
            list_stack.append("bullet" if kind == "bullet_list_open" else "number")
            index += 1
            continue

        if kind in ("bullet_list_close", "ordered_list_close"):
            if list_stack:
                list_stack.pop()
            index += 1
            continue

        if kind == "blockquote_open":
            in_quote += 1
            index += 1
            continue

        if kind == "blockquote_close":
            in_quote = max(0, in_quote - 1)
            index += 1
            continue

        if kind in ("fence", "code_block"):
            _add_code_block(document, token.content)
            index += 1
            continue

        if kind == "hr":
            _add_horizontal_rule(document)
            index += 1
            continue

        if kind == "table_open":
            index = _add_table(document, tokens, index)
            continue

        index += 1


def _new_paragraph(document, list_stack: list[str], in_quote: int):
    if list_stack:
        base = "List Bullet" if list_stack[-1] == "bullet" else "List Number"
        level = min(len(list_stack), _MAX_LIST_STYLE_LEVEL)
        style = base if level == 1 else f"{base} {level}"
        return document.add_paragraph(style=style)
    if in_quote:
        return document.add_paragraph(style="Quote")
    return document.add_paragraph()


def _add_code_block(document, content: str) -> None:
    from docx.shared import Pt

    paragraph = document.add_paragraph(style="No Spacing")
    lines = content.rstrip("\n").split("\n")
    for line_number, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        if line_number < len(lines) - 1:
            run.add_break()


def _add_horizontal_rule(document) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    borders.append(bottom)
    p_pr.append(borders)


def _add_table(document, tokens, start: int) -> int:
    """Render a Markdown table starting at ``table_open`` (index ``start``).

    Returns the index just past the matching ``table_close``."""
    rows: list[list] = []  # each cell is the inline token (or None)
    index = start + 1
    count = len(tokens)
    current: list | None = None

    while index < count and tokens[index].type != "table_close":
        kind = tokens[index].type
        if kind == "tr_open":
            current = []
        elif kind == "tr_close":
            if current is not None:
                rows.append(current)
            current = None
        elif kind in ("th_open", "td_open"):
            inline = tokens[index + 1] if tokens[index + 1].type == "inline" else None
            if current is not None:
                current.append(inline)
        index += 1

    if rows:
        columns = max(len(row) for row in rows)
        table = document.add_table(rows=0, cols=columns)
        table.style = "Light Grid Accent 1"
        for row_index, row in enumerate(rows):
            cells = table.add_row().cells
            for col_index in range(columns):
                inline = row[col_index] if col_index < len(row) else None
                paragraph = cells[col_index].paragraphs[0]
                _render_inline(paragraph, inline, bold=row_index == 0)

    return index + 1  # skip table_close


# --------------------------------------------------------------------- inline


def _render_inline(paragraph, inline_token, *, bold: bool = False) -> None:
    """Emit the children of an ``inline`` token as styled runs."""
    if inline_token is None or not getattr(inline_token, "children", None):
        if inline_token is not None and inline_token.content:
            _add_run(paragraph, inline_token.content, bold=bold)
        return

    state = {"bold": bold, "italic": False, "strike": False}
    for child in inline_token.children:
        kind = child.type
        if kind == "text":
            _add_run(paragraph, child.content, **state)
        elif kind == "code_inline":
            _add_run(paragraph, child.content, code=True, **state)
        elif kind == "strong_open":
            state["bold"] = True
        elif kind == "strong_close":
            state["bold"] = bold
        elif kind in ("em_open",):
            state["italic"] = True
        elif kind in ("em_close",):
            state["italic"] = False
        elif kind == "s_open":
            state["strike"] = True
        elif kind == "s_close":
            state["strike"] = False
        elif kind == "softbreak":
            _add_run(paragraph, " ", **state)
        elif kind == "hardbreak":
            paragraph.add_run().add_break()
        elif kind == "image":
            alt = child.content or (child.attrGet("alt") or "")
            if alt:
                _add_run(paragraph, alt, italic=True)
        # link_open/link_close: the link text is rendered by the text children.


def _add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, strike: bool = False, code: bool = False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.strike = strike
    if code:
        run.font.name = "Consolas"
